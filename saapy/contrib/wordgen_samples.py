import contextlib
import glob
import os
import re
from functools import partial
from pathlib import Path

import docx
import pandas as pd
from invoke import task
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as ALIGN
import numpy as np
from babel.numbers import format_currency, format_decimal
from toolz import compose

# plotting
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns


@contextlib.contextmanager
def directory(dirname=None):
    """
    changes current directory to dirname
    used as with directory(dirname) as dir to
    move out of the changed directory at the end
    :param dirname: directory to change into
    """
    curdir = os.getcwd()
    try:
        if dirname is not None:
            os.chdir(dirname)
        yield Path(dirname or curdir)
    finally:
        os.chdir(curdir)


@task
def install_tools_osx(ctx):
    """
    installs third party tools required to generate diagrams and documents
    from sources including imagemagick, graphviz, pantuml, and pandoc
    :param ctx:
    """
    ctx.run('brew install imagemagick '
            '--with-librsvg --with-pango --with-libwmf --with-openjpeg '
            '--with-ghostscript --with-fftw --with-fontconfig')
    ctx.run('brew install graphviz --with-librsvg --with-freetype '
            '--with-bindings --with-pango')
    ctx.run('brew install plantuml')
    ctx.run('brew install pandoc')


@task
def puml(ctx):
    ctx.run('plantuml -tsvg ./docs/diagrams.puml')


@task
def raster(ctx):
    for svg_path in glob.iglob('./docs/*.svg'):
        cmd = build_raster_command(svg_path)
        print('will run now')
        print(cmd)
        ctx.run(cmd)


def build_raster_command(input_file):
    output_file = re.sub(r'\.svg$', '.png', input_file, 1)
    converter = 'rsvg-convert'
    dpi = 300  # lossless density
    width_in = 6.3  # MS Word A4 page size sans fields
    width_px = int(dpi * width_in)
    size_opts = '--dpi-x={} --width={}'.format(dpi, width_px)
    other_opts = '--format=png --keep-aspect-ratio --background-color=none'
    output_opt = '--output {}'.format(output_file)
    cmd = ' '.join((converter, size_opts, other_opts, output_opt, input_file))
    return cmd


@task(puml, raster)
def dg(ctx):
    pass


@task
def pl(ctx):
    plt.plot([0, 1, 2, 3, 4], [0, 3, 5, 9, 11])
    plt.xlabel('Months')
    plt.ylabel('Books Read')
    plt.savefig('./docs/books_read.png', bbox_inches='tight', dpi=300)


@task
def doc(ctx):
    input_file = 'adding-ds-to-rpa.md'
    output_file = 'adding-ds-to-rpa.docx'
    reference_file = 'style-reference.docx'
    with directory('./docs'):
        ctx.run('pandoc {} -o {} --toc --reference-docx={}'.format(
            input_file, output_file, reference_file))


@task
def doctable(ctx):
    df = pd.read_csv('./docs/flight-options.csv')

    # open an existing document
    doc = docx.Document('./docs/style-reference.docx')

    as_int = partial(format_decimal, format='#')
    as_usd = partial(format_currency, currency='USD')

    s = doc.sections[0]
    width = s.page_width - s.left_margin - s.right_margin

    doc.add_picture('./docs/diagrams_002.png', width=width)

    formatters = {
        'ticket_price': as_usd,
        'total_hours': as_int,
        'trip': as_int,
        'airline': partial(shorten_long_name, width=20),
        'selected': compose({0: 'No', 1: 'Yes'}.get, int)
    }
    add_table(df, doc, table_style='Plain Table 3', formatters=formatters)

    # save the doc
    doc.save('./docs/test.docx')


def add_table(df, word_doc,
              table_style=None,
              include_index=True,
              formatters=None):
    if not formatters:
        formatters = dict()

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row

    first_table_col = 1 if include_index else 0

    table_rows = df.shape[0] + 1
    table_cols = df.shape[1] + first_table_col
    word_table = word_doc.add_table(table_rows,
                                    table_cols,
                                    style=table_style)
    # add the header rows.
    if include_index:
        word_table.cell(0, 0).text = df.index.name or '#'
        index_alignment = prefer_alignment(df.index.dtype)
    alignments = []
    for df_col, table_col in enumerate(range(first_table_col, table_cols)):
        column = df.columns[df_col]
        cell_value = column.replace('_', ' ')
        word_table.cell(0, table_col).text = cell_value
        alignments.append(prefer_alignment(df.dtypes[df_col]))

    # add the rest of the data frame
    for df_row, table_row in enumerate(range(1, table_rows)):
        if include_index:
            cell_value = format_cell_value(
                df.index[df_row],
                formatters,
                column_name='index')
            p = word_table.cell(table_row, 0).paragraphs[0]
            update_paragraph(p, cell_value, index_alignment)
        for df_col, table_col in enumerate(range(first_table_col, table_cols)):
            cell_value = format_cell_value(df.values[df_row, df_col],
                                           formatters,
                                           column_name=df.columns[df_col])
            p = word_table.cell(table_row, table_col).paragraphs[0]
            update_paragraph(p, cell_value, alignments[df_col])

    return word_table


def update_paragraph(p, value, alignment):
    p.text = value
    p.alignment = alignment


def prefer_alignment(value_type):
    if np.issubdtype(value_type, np.number):
        return ALIGN.RIGHT
    else:
        return ALIGN.LEFT


def format_cell_value(value, formatters, column_index=0, column_name=None):
    if column_name and column_name in formatters:
        return formatters[column_name](value)
    elif column_index in formatters:
        return formatters[column_index](value)
    else:
        return str(value)


def shorten_long_name(s, width=50, sep='/', placeholder='...'):
    length = len(s)
    if width <= 0:
        return ''
    elif length <= width:
        return s
    elif length <= len(placeholder):
        return s
    rm_length = length - width + len(placeholder)
    if length <= rm_length:
        return ''
    last_sep_index = s.rfind(sep) if sep else -1
    mid_length = length // 2
    if last_sep_index >= mid_length + rm_length // 2:
        collapse_start = mid_length - rm_length // 2
    elif last_sep_index >= rm_length:
        collapse_start = last_sep_index - rm_length
    else:
        collapse_start = mid_length - rm_length // 2
    collapse_end = collapse_start + rm_length
    collapsed_s = s[:collapse_start] + placeholder + s[collapse_end:]
    return collapsed_s


@task
def genopts(ctx):
    n_trips = 1000
